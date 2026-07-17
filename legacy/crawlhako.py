import os
import re
import asyncio
import aiohttp
import tempfile
import random
import html
import zipfile
import uuid
from datetime import datetime
from urllib.parse import urlparse
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from playwright.async_api import async_playwright
from playwright_stealth import Stealth
from dotenv import load_dotenv

# Sửa lỗi Playwright trên Windows
if os.name == 'nt':
    asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())

# Tải credentials từ .env
load_dotenv()

BASE_URL = 'https://docln.sbs'

# --- CẤU HÌNH ĐĂNG NHẬP ---
LOGIN_URL = os.getenv('LOGIN_URL', 'https://docln.sbs/login')
DOCLN_USERNAME = os.getenv('DOCLN_USERNAME')
DOCLN_PASSWORD = os.getenv('DOCLN_PASSWORD')

STYLE_CSS = '''
@page { margin: 5pt; }
body {
    font-family: "Times New Roman", serif;
    line-height: 1.6;
    text-align: justify;
    margin: 10px;
}
h1 { text-align: center; color: #2c3e50; margin-top: 50px; }
.chapter-header {
    font-size: 20px;
    text-align: center;
    color: #fff;
    padding-bottom: 10px;
    margin-bottom: 30px;
}
h3 { text-align: center; color: #7f8c8d; }
p { margin: 0; text-indent: 1.5em; margin-bottom: 0.5em; }
img {
    display: block;
    margin: 20px auto;
    max-width: 100%;
    height: auto;
    border-radius: 5px;
    box-shadow: 0 4px 8px rgba(0,0,0,0.2);
}
.intro-box { text-align: center; }
.info-list { text-align: left; margin: 20px 0; border-left: 3px solid #34495e; padding-left: 15px; }
.summary-box { text-align: left; font-style: italic; background: #f9f9f9; color: #000; padding: 15px; border-radius: 8px; }
'''

class Utils:
    @staticmethod
    def resolve_url(base_url, relative_url):
        if not relative_url: return ""
        if relative_url.startswith('http'): return relative_url
        if relative_url.startswith('//'): return 'https:' + relative_url

        parsed = urlparse(base_url)
        domain = f"{parsed.scheme}://{parsed.netloc}"
        if relative_url.startswith('/'):
            return domain + relative_url
        return domain + "/" + relative_url

    @staticmethod
    def format_filename(name):
        invalid_chars = ['\\', '/', ':', '*', '?', '"', '<', '>', '|']
        for c in invalid_chars:
            name = name.replace(c, '')
        return name.strip()[:180]

    @staticmethod
    def html_to_plain_text(content_html):
        soup = BeautifulSoup(content_html, "html.parser")
        lines = []

        for element in soup.find_all(['p', 'img']):
            if element.name == 'p':
                text = element.get_text(" ", strip=True)
                if text:
                    lines.append(text)
            elif element.name == 'img':
                img_src = element.get('src', '').split('/')[-1]
                if img_src:
                    lines.append(f"[Hình ảnh: {img_src}]")

        return "\n\n".join(lines).strip()

    @staticmethod
    def html_fragment_to_xhtml(content_html):
        soup = BeautifulSoup(content_html, "html.parser")
        for img in soup.find_all('img'):
            src = img.get('src', '')
            if src.startswith('images/'):
                img['src'] = f"../{src}"
            if not img.get('alt'):
                img['alt'] = img.get('src', '').split('/')[-1] or 'image'
        body = ''.join(str(node) for node in soup.contents).strip()
        return body or '<p></p>'

    @staticmethod
    def summary_to_text(summary_html):
        soup = BeautifulSoup(summary_html or '', "html.parser")
        text = soup.get_text("\n", strip=True)
        return text or "Không có tóm tắt."


def parse_export_formats(raw_formats):
    supported_formats = {'txt', 'docx', 'epub'}
    choice = (raw_formats or '').strip().lower()

    if not choice:
        return None
    if choice == 'all':
        return ['txt', 'docx', 'epub']

    parsed = []
    for item in choice.split(','):
        fmt = item.strip().lower()
        if not fmt:
            continue
        if fmt not in supported_formats:
            return None
        if fmt not in parsed:
            parsed.append(fmt)

    return parsed or None


def export_volume_txt(vol_info, novel_data, chapters_data, save_dir):
    txt_filename = Utils.format_filename(f"{novel_data['title']} - {vol_info['title']}") + ".txt"
    txt_path = os.path.join(save_dir, txt_filename)

    lines = [
        novel_data['title'],
        vol_info['title'],
        f"Tác giả: {novel_data['author']}",
        f"Thể loại: {novel_data['genres']}",
        "",
        "Tóm tắt:",
        Utils.summary_to_text(novel_data['summary']),
        "",
        "=" * 80,
        ""
    ]

    for chap in chapters_data:
        lines.extend([
            chap['title'],
            '-' * max(20, len(chap['title'])),
            Utils.html_to_plain_text(chap['html']),
            "",
            ""
        ])

    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(lines).strip() + "\n")

    print(f"[+] ĐÃ TẠO XONG: {txt_path}")


def export_volume_docx(temp_dir, vol_info, novel_data, chapters_data, save_dir):
    doc = Document()

    title_para = doc.add_heading(novel_data['title'], 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading(vol_info['title'], 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Tác giả: {novel_data['author']}")
    doc.add_paragraph(f"Thể loại: {novel_data['genres']}")
    doc.add_paragraph(f"Tóm tắt: {Utils.summary_to_text(novel_data['summary'])}")

    doc.add_page_break()

    for chap in chapters_data:
        doc.add_heading(chap['title'], level=2)
        soup = BeautifulSoup(chap['html'], "html.parser")

        for element in soup.find_all(['p', 'img']):
            if element.name == 'p':
                text = element.get_text().strip()
                if text:
                    doc.add_paragraph(text)
            elif element.name == 'img':
                img_filename = element.get('src', '').split('/')[-1]
                img_path = os.path.join(temp_dir, "images", img_filename)
                if os.path.exists(img_path):
                    try:
                        doc.add_picture(img_path, width=Inches(5))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        print(f"[-] Error adding image {img_filename} to DOCX: {e}")
                else:
                    print(f"[-] Image not found: {img_path}")

        doc.add_page_break()

    docx_filename = Utils.format_filename(f"{novel_data['title']} - {vol_info['title']}") + ".docx"
    docx_path = os.path.join(save_dir, docx_filename)
    doc.save(docx_path)
    print(f"[+] ĐÃ TẠO XONG: {docx_path}")


def export_volume_epub(temp_dir, vol_info, novel_data, chapters_data, save_dir):
    epub_filename = Utils.format_filename(f"{novel_data['title']} - {vol_info['title']}") + ".epub"
    epub_path = os.path.join(save_dir, epub_filename)
    book_id = str(uuid.uuid4())
    summary_html = Utils.html_fragment_to_xhtml(novel_data['summary'])
    created_at = datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%SZ')

    manifest_items = [
        '<item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>',
        '<item id="ncx" href="toc.ncx" media-type="application/x-dtbncx+xml"/>',
        '<item id="style" href="styles/style.css" media-type="text/css"/>'
    ]
    spine_items = ['<itemref idref="intro"/>']
    nav_points = []
    nav_links = ['<li><a href="text/intro.xhtml">Giới thiệu</a></li>']

    image_items = []
    chapter_files = []

    intro_xhtml = f'''<?xml version="1.0" encoding="utf-8"?>
<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="vi">
  <head>
    <title>{html.escape(novel_data['title'])} - {html.escape(vol_info['title'])}</title>
    <link rel="stylesheet" type="text/css" href="../styles/style.css"/>
  </head>
  <body>
    <h1>{html.escape(novel_data['title'])}</h1>
    <h3>{html.escape(vol_info['title'])}</h3>
    <div class="intro-box">
      <p><strong>Tác giả:</strong> {html.escape(novel_data['author'])}</p>
      <p><strong>Thể loại:</strong> {html.escape(novel_data['genres'])}</p>
    </div>
    <div class="summary-box">
      {summary_html}
    </div>
  </body>
</html>'''

    manifest_items.append('<item id="intro" href="text/intro.xhtml" media-type="application/xhtml+xml"/>')
    nav_points.append('''<navPoint id="navPoint-intro" playOrder="1"><navLabel><text>Giới thiệu</text></navLabel><content src="text/intro.xhtml"/></navPoint>''')

    play_order = 2
    for idx, chap in enumerate(chapters_data, start=1):
        chapter_id = f"chapter_{idx}"
        chapter_file = f"text/chapter_{idx:03d}.xhtml"
        chapter_title = chap['title'] or f"Chương {idx}"
        chapter_body = Utils.html_fragment_to_xhtml(chap['html'])

        chapter_xhtml = f'''<?xml version="1.0" encoding="utf-8"?>
<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="vi">
  <head>
    <title>{html.escape(chapter_title)}</title>
    <link rel="stylesheet" type="text/css" href="../styles/style.css"/>
  </head>
  <body>
    <div class="chapter-header">{html.escape(chapter_title)}</div>
    {chapter_body}
  </body>
</html>'''

        chapter_files.append((chapter_file, chapter_xhtml))
        manifest_items.append(f'<item id="{chapter_id}" href="{chapter_file}" media-type="application/xhtml+xml"/>')
        spine_items.append(f'<itemref idref="{chapter_id}"/>')
        nav_points.append(
            f'<navPoint id="navPoint-{idx}" playOrder="{play_order}"><navLabel><text>{html.escape(chapter_title)}</text></navLabel><content src="{chapter_file}"/></navPoint>'
        )
        nav_links.append(f'<li><a href="{chapter_file}">{html.escape(chapter_title)}</a></li>')
        play_order += 1

    image_id = 1
    added_images = set()
    for chap in chapters_data:
        for img_name in chap.get('images', []):
            if img_name in added_images:
                continue
            added_images.add(img_name)
            ext = os.path.splitext(img_name)[1].lower()
            media_type = {
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.png': 'image/png',
                '.gif': 'image/gif',
                '.webp': 'image/webp',
                '.svg': 'image/svg+xml'
            }.get(ext, 'image/jpeg')
            image_items.append((img_name, os.path.join(temp_dir, 'images', img_name), media_type, f'image_{image_id}'))
            manifest_items.append(f'<item id="image_{image_id}" href="images/{img_name}" media-type="{media_type}"/>')
            image_id += 1

    content_opf = f'''<?xml version="1.0" encoding="utf-8"?>
<package xmlns="http://www.idpf.org/2007/opf" unique-identifier="BookId" version="2.0">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:opf="http://www.idpf.org/2007/opf">
    <dc:title>{html.escape(novel_data['title'])} - {html.escape(vol_info['title'])}</dc:title>
    <dc:creator>{html.escape(novel_data['author'])}</dc:creator>
    <dc:language>vi</dc:language>
    <dc:identifier id="BookId">urn:uuid:{book_id}</dc:identifier>
    <dc:description>{html.escape(Utils.summary_to_text(novel_data['summary']))}</dc:description>
    <dc:date>{created_at}</dc:date>
  </metadata>
  <manifest>
    {' '.join(manifest_items)}
  </manifest>
  <spine toc="ncx">
    {' '.join(spine_items)}
  </spine>
</package>'''

    toc_ncx = f'''<?xml version="1.0" encoding="utf-8"?>
<ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">
  <head>
    <meta name="dtb:uid" content="urn:uuid:{book_id}"/>
    <meta name="dtb:depth" content="1"/>
    <meta name="dtb:totalPageCount" content="0"/>
    <meta name="dtb:maxPageNumber" content="0"/>
  </head>
  <docTitle><text>{html.escape(novel_data['title'])} - {html.escape(vol_info['title'])}</text></docTitle>
  <navMap>
    {''.join(nav_points)}
  </navMap>
</ncx>'''

    nav_xhtml = f'''<?xml version="1.0" encoding="utf-8"?>
<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="vi">
  <head>
    <title>Mục lục</title>
    <link rel="stylesheet" type="text/css" href="styles/style.css"/>
  </head>
  <body>
    <h1>Mục lục</h1>
    <ol>
      {''.join(nav_links)}
    </ol>
  </body>
</html>'''

    os.makedirs(save_dir, exist_ok=True)
    with zipfile.ZipFile(epub_path, 'w') as epub:
        epub.writestr('mimetype', 'application/epub+zip', compress_type=zipfile.ZIP_STORED)
        epub.writestr('META-INF/container.xml', '''<?xml version="1.0"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles>
    <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>
  </rootfiles>
</container>''')
        epub.writestr('OEBPS/content.opf', content_opf)
        epub.writestr('OEBPS/toc.ncx', toc_ncx)
        epub.writestr('OEBPS/nav.xhtml', nav_xhtml)
        epub.writestr('OEBPS/styles/style.css', STYLE_CSS)
        epub.writestr('OEBPS/text/intro.xhtml', intro_xhtml)

        for chapter_file, chapter_xhtml in chapter_files:
            epub.writestr(f'OEBPS/{chapter_file}', chapter_xhtml)

        for img_name, img_path, _, _ in image_items:
            if os.path.exists(img_path):
                epub.write(img_path, f'OEBPS/images/{img_name}')

    print(f"[+] ĐÃ TẠO XONG: {epub_path}")


async def async_auto_login(page):
    """Tự động đăng nhập vào docln.sbs (async version)"""
    try:
        # Kiểm tra xem đã login chưa
        print(f"🔐 [Login] Đang điều hướng đến: {LOGIN_URL}")
        await page.goto(LOGIN_URL, wait_until="domcontentloaded", timeout=60000)
        await page.wait_for_load_state("networkidle")
        
        # Nếu có credentials từ .env
        if DOCLN_USERNAME and DOCLN_PASSWORD:
            print("🔐 Đang thực hiện đăng nhập tự động...")
            print(f"📝 Đang điền thông tin đăng nhập...")
            
            # Tìm email input
            email_input = None
            
            # Cách 1: Tìm bằng placeholder
            try:
                email_input = page.locator('input[placeholder*="Email"]').first
                if email_input and await email_input.is_visible():
                    print("✓ Tìm thấy email input bằng placeholder")
                else:
                    email_input = None
            except:
                pass
            
            # Cách 2: Tìm input đầu tiên
            if not email_input:
                try:
                    all_inputs = await page.locator('input[type="text"], input:not([type])').all()
                    if len(all_inputs) > 0:
                        email_input = all_inputs[0]
                        print("✓ Tìm thấy email input (input đầu tiên)")
                except:
                    pass
            
            if not email_input:
                print("❌ Không tìm thấy email input")
                raise Exception("Cannot find email input field")
            
            # Điền email
            await email_input.fill(DOCLN_USERNAME)
            print(f"✓ Đã điền email: {DOCLN_USERNAME}")
            
            # Tìm password input
            password_input = page.locator('input[type="password"]').first
            
            if not password_input or not await password_input.is_visible():
                print("❌ Không tìm thấy password input")
                raise Exception("Cannot find password input field")
            
            # Điền password
            await password_input.fill(DOCLN_PASSWORD)
            print("✓ Đã điền mật khẩu")
            
            # Đợi một chút
            await page.wait_for_timeout(500)
            
            # Tìm button submit
            submit_btn = None
            
            # Cách 1: Tìm button có text "Đăng nhập"
            try:
                buttons = await page.locator('button').all()
                for btn in buttons:
                    if await btn.is_visible():
                        text = await btn.text_content()
                        if "Đăng nhập" in text:
                            submit_btn = btn
                            print(f"✓ Tìm thấy button: {text.strip()}")
                            break
            except:
                pass
            
            # Cách 2: Button[type="submit"]
            if not submit_btn:
                try:
                    submit_btn = page.locator('button[type="submit"]').first
                    if submit_btn and await submit_btn.is_visible():
                        print("✓ Tìm thấy button submit")
                except:
                    pass
            
            if not submit_btn:
                print("❌ Không tìm thấy button Đăng nhập")
                raise Exception("Cannot find submit button")
            
            # Click button
            print("✓ Đang click button Đăng nhập...")
            await submit_btn.click()
            
            # Đợi xử lý
            print("⏳ Đợi server xử lý...")
            await page.wait_for_timeout(2000)
            
            # Kiểm tra kết quả
            current_url = page.url
            print(f"📍 URL sau login: {current_url}")
            
            # Nếu vẫn ở trang login
            if "login" in current_url.lower():
                # Check xem có lỗi không
                print("⚠️  Vẫn ở trang login, kiểm tra lỗi...")
                
                # Tìm error messages
                try:
                    error_elements = await page.locator('.alert, .error, [class*="danger"]').all()
                    if error_elements:
                        for error in error_elements:
                            if await error.is_visible():
                                error_text = await error.text_content()
                                print(f"❌ Error: {error_text.strip()}")
                except:
                    pass
                
                print("⚠️  Đăng nhập thất bại - Chuyển sang manual mode")
                input("Hãy đăng nhập thủ công, sau đó nhấn Enter...")
                return True
            else:
                print("✅ Đăng nhập thành công!")
                return True
                
        else:
            print("⚠️  Không tìm thấy DOCLN_USERNAME hoặc DOCLN_PASSWORD trong .env")
            print("📝 Vui lòng đăng nhập thủ công...")
            input("Đã đăng nhập xong? Nhấn Enter để tiếp tục...")
            return True
            
    except Exception as e:
        print(f"⚠️  Lỗi trong quá trình đăng nhập: {e}")
        print("📝 Vui lòng đăng nhập thủ công...")
        print("💡 Gợi ý: Đảm bảo credentials trong .env là đúng")
        input("Đã đăng nhập xong? Nhấn Enter để tiếp tục...")
        return True


class HakoAsyncCrawler:
    def __init__(self, temp_dir):
        self.temp_dir = temp_dir
        self.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        }

    async def fetch_info(self, url):
        print(f"\n[*] Đang khởi động trình duyệt tàng hình (Stealth) để vượt Cloudflare...")
        async with Stealth().use_async(async_playwright()) as p:
            browser = await p.chromium.launch(headless=False)
            context = await browser.new_context()
            page = await context.new_page()

            try:
                # Thực hiện đăng nhập nếu có credentials
                if DOCLN_USERNAME and DOCLN_PASSWORD:
                    print("[*] Phát hiện credentials - Thực hiện đăng nhập...")
                    await async_auto_login(page)
                    print("[*] Tiếp tục tải thông tin truyện...\n")
                
                print(f"[*] Đang truy cập: {url}")
                await page.goto(url, timeout=60000)
                await page.wait_for_selector(".series-name, .volume-list", timeout=15000)

                html = await page.content()
                soup = BeautifulSoup(html, "html.parser")

                title_tag = soup.select_one('.series-name a')
                novel_title = title_tag.text.strip() if title_tag else "Không rõ tên truyện"

                author = "Không rõ"
                for info in soup.select('.info-item'):
                    if "Tác giả:" in info.text: 
                        author = info.select_one('.info-value').text.strip()

                genres = ", ".join([a.text.strip() for a in soup.select('.series-gernes a')])
                summary_div = soup.select_one('.summary-content')
                summary_html = str(summary_div) if summary_div else "Không có tóm tắt."

                series_cover_bytes = None
                series_cover_div = soup.select_one('.series-cover .img-in-ratio')
                if series_cover_div and 'style' in series_cover_div.attrs:
                    s_match = re.search(r"url\(['\"]?(.*?)['\"]?\)", series_cover_div['style'])
                    if s_match:
                        s_img_url = Utils.resolve_url(url, s_match.group(1))
                        # Tắt xác thực SSL ở đây
                        async with aiohttp.ClientSession(connector=aiohttp.TCPConnector(ssl=False)) as session:
                            async with session.get(s_img_url) as resp:
                                if resp.status == 200: series_cover_bytes = await resp.read()

                volumes = []
                # Tắt xác thực SSL ở đây
                async with aiohttp.ClientSession(connector=aiohttp.TCPConnector(ssl=False)) as session:
                    for vol_section in soup.select('section.volume-list'):
                        vol_title_tag = vol_section.select_one('.sect-title')
                        if not vol_title_tag: continue
                        vol_title = vol_title_tag.text.replace('*', '').strip()

                        vol_cover_bytes = None
                        v_cover_div = vol_section.select_one('.volume-cover .img-in-ratio')
                        if v_cover_div and 'style' in v_cover_div.attrs:
                            v_match = re.search(r"url\(['\"]?(.*?)['\"]?\)", v_cover_div['style'])
                            if v_match:
                                v_img_url = Utils.resolve_url(url, v_match.group(1))
                                try:
                                    async with session.get(v_img_url) as resp:
                                        if resp.status == 200: vol_cover_bytes = await resp.read()
                                except: pass

                        chapters = []
                        for a_tag in vol_section.select('.list-chapters li .chapter-name a'):
                            chapters.append({
                                "title": a_tag.text.strip(),
                                "url": Utils.resolve_url(url, a_tag.get('href'))
                            })

                        if chapters:
                            volumes.append({
                                "title": vol_title,
                                "chapters": chapters,
                                "cover_bytes": vol_cover_bytes if vol_cover_bytes else series_cover_bytes
                            })

                print(f"[+] Tìm thấy truyện: {novel_title} ({len(volumes)} tập)")
                return {
                    "title": novel_title, "author": author, 
                    "genres": genres, "summary": summary_html,
                    "cover_bytes": series_cover_bytes, "volumes": volumes
                }

            except Exception as e:
                print(f"[-] Lỗi tải thông tin: {e}")
                return None
            finally:
                await browser.close()

    async def download_chapter(self, page, session, chapter_info, chapter_idx):
        try:
            await page.goto(chapter_info['url'], timeout=60000)
            await page.wait_for_selector("#chapter-content", state="attached", timeout=15000)
            raw_html = await page.locator("#chapter-content").inner_html()

            soup = BeautifulSoup(raw_html, "html.parser")

            for r in soup.select('#chapter-c-protected, p[style*="display: none"]'): r.decompose()
            for banner in soup.find_all('a', href=re.compile(r"/truyen/\d+")):
                if banner.find('img', src=re.compile(r"chapter-banners|banners")): banner.decompose()
            for social in soup.select('a[href*="discord.gg"], a[href*="facebook.com"]'): social.decompose()

            for note in soup.select('.inline-note, .note-content'):
                note.insert_before(soup.new_string(" [Chú thích: "))
                note.insert_after(soup.new_string("] "))
                if 'style' in note.attrs: del note.attrs['style']

            # Tạo thư mục images trong temp_dir
            images_dir = os.path.join(self.temp_dir, "images")
            os.makedirs(images_dir, exist_ok=True)
            
            chapter_images = []
            tasks = []
            
            for i, img_tag in enumerate(soup.find_all('img')):
                raw_img_url = img_tag.get('src')
                if not raw_img_url: continue

                img_url = Utils.resolve_url(chapter_info['url'], raw_img_url)
                if "icon" in img_url or "tracker" in img_url:
                    img_tag.decompose()
                    continue

                img_ext = img_url.split('.')[-1].split('?')[0]
                if len(img_ext) > 4: img_ext = "jpg"
                local_name = f"chap{chapter_idx}_img{i}.{img_ext}"
                save_path = os.path.join(images_dir, local_name)

                img_tag['src'] = f"images/{local_name}"
                chapter_images.append(local_name)
                
                tasks.append(self._download_image(session, img_url, save_path))

            if tasks:
                results = await asyncio.gather(*tasks, return_exceptions=True)
                for idx, result in enumerate(results):
                    if isinstance(result, Exception):
                        print(f"[-] Error downloading image {idx}: {result}")

            clean_content = str(soup).replace('<br>', '<br/>').replace('<hr>', '<hr/>')
            clean_content = re.sub(r'<img([^>]*?)(?<!/)>', r'<img\1/>', clean_content)
            
            return clean_content, chapter_images

        except Exception as e:
            print(f"[-] Lỗi tải chương {chapter_info['title']}: {e}")
            return "<p>Chương này bị lỗi tải nội dung.</p>", []

    async def _download_image(self, session, url, save_path):
        try:
            async with session.get(url, timeout=30) as response:
                if response.status == 200:
                    # Đảm bảo thư mục tồn tại
                    os.makedirs(os.path.dirname(save_path), exist_ok=True)
                    with open(save_path, 'wb') as f:
                        f.write(await response.read())
                    return True
                else:
                    # Use ASCII-safe error message to avoid encoding issues
                    print(f"[-] Error downloading image {url}: HTTP {response.status}")
                    return False
        except Exception as e:
            # Use ASCII-safe error message to avoid encoding issues
            print(f"[-] Error downloading image {url}: {e}")
            return False

async def process_volume(vol_info, novel_data, save_dir, export_formats):
    print(f"\n{'-'*40}\n[*] Đang khởi tạo luồng tải tập: {vol_info['title']}")
    
    with tempfile.TemporaryDirectory() as temp_dir:
        crawler = HakoAsyncCrawler(temp_dir)
        chapters_data = []
        
        async with Stealth().use_async(async_playwright()) as p:
            browser = await p.chromium.launch(headless=False)
            context = await browser.new_context()
            page = await context.new_page()

            # Thực hiện đăng nhập nếu có credentials
            if DOCLN_USERNAME and DOCLN_PASSWORD:
                print("[*] Phát hiện credentials - Thực hiện đăng nhập...")
                await async_auto_login(page)
                print("[*] Tiếp tục tải nội dung các chương...\n")

            async with aiohttp.ClientSession(connector=aiohttp.TCPConnector(ssl=False)) as session:
                chapters = vol_info['chapters']
                for idx, chapter in enumerate(chapters):
                    print(f"[*] Đang tải chương {idx+1}/{len(chapters)}: {chapter['title']}")
                    html, imgs = await crawler.download_chapter(page, session, chapter, idx)
                    chapters_data.append({"title": chapter['title'], "html": html, "images": imgs})

                    if idx < len(chapters) - 1:
                        await asyncio.sleep(random.uniform(2.0, 4.0))

            await browser.close()

        os.makedirs(save_dir, exist_ok=True)

        if 'txt' in export_formats:
            print("[*] Đang đóng gói bản TXT...")
            export_volume_txt(vol_info, novel_data, chapters_data, save_dir)

        if 'docx' in export_formats:
            print("[*] Đang đóng gói bản Word (DOCX)...")
            export_volume_docx(temp_dir, vol_info, novel_data, chapters_data, save_dir)

        if 'epub' in export_formats:
            print("[*] Đang đóng gói bản EPUB...")
            export_volume_epub(temp_dir, vol_info, novel_data, chapters_data, save_dir)

async def main():
    os.system('cls' if os.name == 'nt' else 'clear')
    print("=== DOCLN/HAKO CRAWLER PRO (ASYNC CLI) ===")
    url = input("Nhập URL bộ truyện (VD: https://docln.net/truyen/1234): ").strip()
    if not url: return

    with tempfile.TemporaryDirectory() as tmp:
        crawler = HakoAsyncCrawler(tmp)
        novel_data = await crawler.fetch_info(url)
        
    if not novel_data: return

    print("\n--- DANH SÁCH CÁC TẬP ---")
    for i, vol in enumerate(novel_data['volumes']):
        print(f"[{i}] {vol['title']}")

    choices = input("\nNhập số thứ tự các tập cần tải (cách nhau bằng dấu phẩy, VD: 0,1,2 | Hoặc gõ 'all'): ").strip()
    
    selected_indices = []
    if choices.lower() == 'all':
        selected_indices = list(range(len(novel_data['volumes'])))
    else:
        try:
            selected_indices = [int(x.strip()) for x in choices.split(',') if x.strip().isdigit()]
        except ValueError:
            print("[-] Lựa chọn không hợp lệ.")
            return

    save_dir = Utils.format_filename(novel_data['title'])
    format_choices = input("Nhập định dạng cần xuất (txt, docx, epub | nhiều định dạng cách nhau bằng dấu phẩy | hoặc gõ 'all'): ").strip()
    export_formats = parse_export_formats(format_choices)

    if not export_formats:
        print("[-] Định dạng không hợp lệ. Chỉ hỗ trợ: txt, docx, epub hoặc all.")
        return
    
    for idx in selected_indices:
        if 0 <= idx < len(novel_data['volumes']):
            await process_volume(novel_data['volumes'][idx], novel_data, save_dir, export_formats)
        else:
            print(f"[-] Bỏ qua chỉ số không hợp lệ: {idx}")

    print(f"\n[+] HOÀN TẤT TẤT CẢ! Thư mục lưu: ./{save_dir}")

if __name__ == "__main__":
    import sys
    import asyncio

    if sys.platform == 'win32':
        asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())
        
    asyncio.run(main())