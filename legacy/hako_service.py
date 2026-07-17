
import asyncio
import html
import os
import random
import re
import tempfile
import uuid
import zipfile
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Callable, Dict, List, Optional, Tuple
from urllib.parse import urlparse

import aiohttp
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from playwright.async_api import Browser, BrowserContext, Page, async_playwright
from playwright_stealth import Stealth

BASE_URL = "https://docln.sbs"
LOGIN_URL = f"{BASE_URL}/login"
STYLE_CSS = """
@page { margin: 5pt; }
body {
    font-family: "Times New Roman", serif;
    line-height: 1.6;
    text-align: justify;
    margin: 10px;
}
h1 { text-align: center; color: #2c3e50; margin-top: 50px; }
.chapter-header {
    font-size: 20px;
    text-align: center;
    color: #fff;
    padding-bottom: 10px;
    margin-bottom: 30px;
}
h3 { text-align: center; color: #7f8c8d; }
p { margin: 0; text-indent: 1.5em; margin-bottom: 0.5em; }
img {
    display: block;
    margin: 20px auto;
    max-width: 100%;
    height: auto;
    border-radius: 5px;
    box-shadow: 0 4px 8px rgba(0,0,0,0.2);
}
.intro-box { text-align: center; }
.info-list { text-align: left; margin: 20px 0; border-left: 3px solid #34495e; padding-left: 15px; }
.summary-box { text-align: left; font-style: italic; background: #f9f9f9; color: #000; padding: 15px; border-radius: 8px; }
"""


class HakoError(Exception):
    pass


class AuthenticationError(HakoError):
    pass


class Utils:
    @staticmethod
    def resolve_url(base_url: str, relative_url: str) -> str:
        if not relative_url:
            return ""
        if relative_url.startswith("http"):
            return relative_url
        if relative_url.startswith("//"):
            return "https:" + relative_url
        parsed = urlparse(base_url)
        domain = f"{parsed.scheme}://{parsed.netloc}"
        if relative_url.startswith("/"):
            return domain + relative_url
        return domain + "/" + relative_url

    @staticmethod
    def format_filename(name: str) -> str:
        for c in ['\\', '/', ':', '*', '?', '"', '<', '>', '|']:
            name = name.replace(c, '')
        return name.strip()[:180] or "untitled"

    @staticmethod
    def summary_to_text(summary_html: str) -> str:
        soup = BeautifulSoup(summary_html or "", "html.parser")
        text = soup.get_text("\n", strip=True)
        return text or "Không có tóm tắt."

    @staticmethod
    def html_to_plain_text(content_html: str) -> str:
        soup = BeautifulSoup(content_html, "html.parser")
        lines = []
        for element in soup.find_all(['p', 'img']):
            if element.name == 'p':
                text = element.get_text(" ", strip=True)
                if text:
                    lines.append(text)
            elif element.name == 'img':
                img_src = element.get('src', '').split('/')[-1]
                if img_src:
                    lines.append(f"[Hình ảnh: {img_src}]")
        return "\n\n".join(lines).strip()

    @staticmethod
    def html_fragment_to_xhtml(content_html: str) -> str:
        soup = BeautifulSoup(content_html, "html.parser")
        for img in soup.find_all('img'):
            src = img.get('src', '')
            if src.startswith('images/'):
                img['src'] = f"../{src}"
            if not img.get('alt'):
                img['alt'] = img.get('src', '').split('/')[-1] or 'image'
        body = ''.join(str(node) for node in soup.contents).strip()
        return body or '<p></p>'


@dataclass
class SessionState:
    browser: Optional[Browser] = None
    context: Optional[BrowserContext] = None
    page: Optional[Page] = None
    playwright = None
    lock: asyncio.Lock = field(default_factory=asyncio.Lock)
    logged_in: bool = False
    user_label: str = "Chưa đăng nhập"


class HakoSessionManager:
    def __init__(self, state: SessionState):
        self.state = state

    async def ensure_browser(self):
        async with self.state.lock:
            if self.state.browser and self.state.context and self.state.page:
                return
            self.state.playwright = await async_playwright().start()
            self.state.browser = await self.state.playwright.chromium.launch(headless=True)
            self.state.context = await self.state.browser.new_context()
            await Stealth().apply_stealth_async(self.state.context)
            self.state.page = await self.state.context.new_page()

    async def login(self, username: str, password: str) -> str:
        await self.ensure_browser()
        page = self.state.page
        assert page is not None

        await page.goto(LOGIN_URL, wait_until="domcontentloaded", timeout=60000)
        await page.wait_for_load_state("networkidle")

        email_input = page.locator('input[placeholder*="Email"], input[type="email"], input[type="text"]').first
        password_input = page.locator('input[type="password"]').first
        submit_button = page.locator('button[type="submit"], button:has-text("Đăng nhập"), button:has-text("Dang nhap")').first

        await email_input.fill(username)
        await password_input.fill(password)
        await submit_button.click()
        await page.wait_for_timeout(2500)
        await page.wait_for_load_state("networkidle")

        current_url = page.url.lower()
        content = await page.content()
        if "login" in current_url:
            soup = BeautifulSoup(content, "html.parser")
            alert = soup.select_one('.alert, .error, .validation-summary-errors, [class*="danger"]')
            detail = alert.get_text(" ", strip=True) if alert else "Không thể đăng nhập. Hãy kiểm tra tài khoản hoặc captcha/anti-bot."
            self.state.logged_in = False
            raise AuthenticationError(detail)

        label = await self._read_user_label(page)
        self.state.logged_in = True
        self.state.user_label = label or username
        return self.state.user_label

    async def _read_user_label(self, page: Page) -> str:
        selectors = [
            '.header-user-name',
            '.dropdown-user strong',
            '.user-name',
            'a[href*="/thanh-vien/"]',
            'a[href*="/member/"]'
        ]
        for selector in selectors:
            try:
                locator = page.locator(selector).first
                if await locator.count() and await locator.is_visible():
                    text = (await locator.text_content() or '').strip()
                    if text:
                        return text
            except Exception:
                continue
        return "Đăng nhập thành công"

    async def export_storage_state(self) -> Dict:
        await self.ensure_browser()
        assert self.state.context is not None
        return await self.state.context.storage_state()

    async def close(self):
        async with self.state.lock:
            if self.state.context:
                await self.state.context.close()
            if self.state.browser:
                await self.state.browser.close()
            if self.state.playwright:
                await self.state.playwright.stop()
            self.state.browser = None
            self.state.context = None
            self.state.page = None
            self.state.playwright = None
            self.state.logged_in = False
            self.state.user_label = "Chưa đăng nhập"


class HakoCrawler:
    def __init__(self, storage_state: Optional[Dict] = None, progress_cb: Optional[Callable[[str], None]] = None):
        self.storage_state = storage_state
        self.progress_cb = progress_cb or (lambda _msg: None)

    def log(self, message: str):
        self.progress_cb(message)

    async def _new_page(self) -> Tuple:
        playwright = await async_playwright().start()
        browser = await playwright.chromium.launch(headless=True)
        context = await browser.new_context(storage_state=self.storage_state)
        await Stealth().apply_stealth_async(context)
        page = await context.new_page()
        return playwright, browser, context, page

    async def fetch_info(self, url: str) -> Dict:
        playwright, browser, context, page = await self._new_page()
        try:
            self.log(f"Đang truy cập bộ truyện: {url}")
            await page.goto(url, timeout=60000)
            await page.wait_for_selector('.series-name, .volume-list', timeout=20000)
            html_doc = await page.content()
            soup = BeautifulSoup(html_doc, 'html.parser')

            title_tag = soup.select_one('.series-name a, .series-name')
            novel_title = title_tag.get_text(strip=True) if title_tag else 'Không rõ tên truyện'

            author = 'Không rõ'
            for info in soup.select('.info-item'):
                text = info.get_text(' ', strip=True)
                if 'Tác giả' in text:
                    value = info.select_one('.info-value')
                    author = value.get_text(strip=True) if value else text
                    break

            genres = ', '.join([a.get_text(strip=True) for a in soup.select('.series-gernes a, .series-genres a')]) or 'Không rõ'
            summary_div = soup.select_one('.summary-content')
            summary_html = str(summary_div) if summary_div else '<p>Không có tóm tắt.</p>'

            volumes = []
            for vol_index, vol_section in enumerate(soup.select('section.volume-list'), start=1):
                vol_title_tag = vol_section.select_one('.sect-title')
                if not vol_title_tag:
                    continue
                vol_title = vol_title_tag.get_text(' ', strip=True).replace('*', '')
                chapters = []
                for a_tag in vol_section.select('.list-chapters li .chapter-name a'):
                    chapters.append({
                        'title': a_tag.get_text(' ', strip=True),
                        'url': Utils.resolve_url(url, a_tag.get('href')),
                    })
                if chapters:
                    volumes.append({
                        'id': vol_index - 1,
                        'title': vol_title,
                        'chapter_count': len(chapters),
                        'chapters': chapters,
                    })

            return {
                'title': novel_title,
                'author': author,
                'genres': genres,
                'summary': summary_html,
                'volumes': volumes,
                'source_url': url,
            }
        finally:
            await context.close()
            await browser.close()
            await playwright.stop()

    async def download_chapter(self, page: Page, session: aiohttp.ClientSession, temp_dir: str, chapter_info: Dict, chapter_idx: int):
        try:
            await page.goto(chapter_info['url'], timeout=60000)
            await page.wait_for_selector('#chapter-content', state='attached', timeout=15000)
            raw_html = await page.locator('#chapter-content').inner_html()
            soup = BeautifulSoup(raw_html, 'html.parser')

            for r in soup.select('#chapter-c-protected, p[style*="display: none"]'):
                r.decompose()
            for banner in soup.find_all('a', href=re.compile(r'/truyen/\d+')):
                if banner.find('img', src=re.compile(r'chapter-banners|banners')):
                    banner.decompose()
            for social in soup.select('a[href*="discord.gg"], a[href*="facebook.com"]'):
                social.decompose()
            for note in soup.select('.inline-note, .note-content'):
                note.insert_before(soup.new_string(' [Chú thích: '))
                note.insert_after(soup.new_string('] '))
                if 'style' in note.attrs:
                    del note.attrs['style']

            images_dir = os.path.join(temp_dir, 'images')
            os.makedirs(images_dir, exist_ok=True)

            chapter_images = []
            tasks = []
            for i, img_tag in enumerate(soup.find_all('img')):
                raw_img_url = img_tag.get('src')
                if not raw_img_url:
                    continue
                img_url = Utils.resolve_url(chapter_info['url'], raw_img_url)
                if 'icon' in img_url or 'tracker' in img_url:
                    img_tag.decompose()
                    continue

                img_ext = img_url.split('.')[-1].split('?')[0].lower()
                if len(img_ext) > 5 or '/' in img_ext:
                    img_ext = 'jpg'
                local_name = f"chap{chapter_idx}_img{i}.{img_ext}"
                save_path = os.path.join(images_dir, local_name)
                img_tag['src'] = f"images/{local_name}"
                chapter_images.append(local_name)
                tasks.append(self._download_image(session, img_url, save_path))

            if tasks:
                await asyncio.gather(*tasks, return_exceptions=True)

            clean_content = str(soup).replace('<br>', '<br/>').replace('<hr>', '<hr/>')
            clean_content = re.sub(r'<img([^>]*?)(?<!/)>', r'<img\1/>', clean_content)
            return clean_content, chapter_images
        except Exception as exc:
            self.log(f"Lỗi tải chương {chapter_info['title']}: {exc}")
            return '<p>Chương này bị lỗi tải nội dung.</p>', []

    async def _download_image(self, session: aiohttp.ClientSession, url: str, save_path: str):
        try:
            async with session.get(url, timeout=30) as response:
                if response.status == 200:
                    os.makedirs(os.path.dirname(save_path), exist_ok=True)
                    with open(save_path, 'wb') as f:
                        f.write(await response.read())
        except Exception:
            return False
        return True


def parse_export_formats(raw_formats: List[str]) -> List[str]:
    supported = ['txt', 'docx', 'epub']
    result = [fmt for fmt in raw_formats if fmt in supported]
    return result or ['epub']


def export_volume_txt(vol_info, novel_data, chapters_data, save_dir):
    txt_filename = Utils.format_filename(f"{novel_data['title']} - {vol_info['title']}") + '.txt'
    txt_path = os.path.join(save_dir, txt_filename)
    lines = [
        novel_data['title'],
        vol_info['title'],
        f"Tác giả: {novel_data['author']}",
        f"Thể loại: {novel_data['genres']}",
        '',
        'Tóm tắt:',
        Utils.summary_to_text(novel_data['summary']),
        '',
        '=' * 80,
        ''
    ]
    for chap in chapters_data:
        lines.extend([
            chap['title'],
            '-' * max(20, len(chap['title'])),
            Utils.html_to_plain_text(chap['html']),
            '',
            ''
        ])
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(lines).strip() + "\n")
    return txt_path


def export_volume_docx(temp_dir, vol_info, novel_data, chapters_data, save_dir):
    doc = Document()
    title_para = doc.add_heading(novel_data['title'], 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(vol_info['title'], 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Tác giả: {novel_data['author']}")
    doc.add_paragraph(f"Thể loại: {novel_data['genres']}")
    doc.add_paragraph(f"Tóm tắt: {Utils.summary_to_text(novel_data['summary'])}")
    doc.add_page_break()

    for chap in chapters_data:
        doc.add_heading(chap['title'], level=2)
        soup = BeautifulSoup(chap['html'], 'html.parser')
        for element in soup.find_all(['p', 'img']):
            if element.name == 'p':
                text = element.get_text().strip()
                if text:
                    doc.add_paragraph(text)
            elif element.name == 'img':
                img_filename = element.get('src', '').split('/')[-1]
                img_path = os.path.join(temp_dir, 'images', img_filename)
                if os.path.exists(img_path):
                    try:
                        doc.add_picture(img_path, width=Inches(5))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        pass
        doc.add_page_break()

    docx_filename = Utils.format_filename(f"{novel_data['title']} - {vol_info['title']}") + '.docx'
    docx_path = os.path.join(save_dir, docx_filename)
    doc.save(docx_path)
    return docx_path


def export_volume_epub(temp_dir, vol_info, novel_data, chapters_data, save_dir):
    epub_filename = Utils.format_filename(f"{novel_data['title']} - {vol_info['title']}") + '.epub'
    epub_path = os.path.join(save_dir, epub_filename)
    book_id = str(uuid.uuid4())
    summary_html = Utils.html_fragment_to_xhtml(novel_data['summary'])
    created_at = datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%SZ')

    manifest_items = [
        '<item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>',
        '<item id="ncx" href="toc.ncx" media-type="application/x-dtbncx+xml"/>',
        '<item id="style" href="styles/style.css" media-type="text/css"/>'
    ]
    spine_items = ['<itemref idref="intro"/>']
    nav_points = []
    nav_links = ['<li><a href="text/intro.xhtml">Giới thiệu</a></li>']
    chapter_files = []

    intro_xhtml = f"""<?xml version="1.0" encoding="utf-8"?>
<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="vi">
  <head>
    <title>{html.escape(novel_data['title'])} - {html.escape(vol_info['title'])}</title>
    <link rel="stylesheet" type="text/css" href="../styles/style.css"/>
  </head>
  <body>
    <h1>{html.escape(novel_data['title'])}</h1>
    <h3>{html.escape(vol_info['title'])}</h3>
    <div class="intro-box">
      <p><strong>Tác giả:</strong> {html.escape(novel_data['author'])}</p>
      <p><strong>Thể loại:</strong> {html.escape(novel_data['genres'])}</p>
    </div>
    <div class="summary-box">{summary_html}</div>
  </body>
</html>"""
    manifest_items.append('<item id="intro" href="text/intro.xhtml" media-type="application/xhtml+xml"/>')
    nav_points.append('<navPoint id="navPoint-intro" playOrder="1"><navLabel><text>Giới thiệu</text></navLabel><content src="text/intro.xhtml"/></navPoint>')

    play_order = 2
    for idx, chap in enumerate(chapters_data, start=1):
        chapter_id = f"chapter_{idx}"
        chapter_file = f"text/chapter_{idx:03d}.xhtml"
        chapter_title = chap['title'] or f"Chương {idx}"
        chapter_body = Utils.html_fragment_to_xhtml(chap['html'])
        chapter_xhtml = f"""<?xml version="1.0" encoding="utf-8"?>
<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="vi">
  <head>
    <title>{html.escape(chapter_title)}</title>
    <link rel="stylesheet" type="text/css" href="../styles/style.css"/>
  </head>
  <body>
    <div class="chapter-header">{html.escape(chapter_title)}</div>
    {chapter_body}
  </body>
</html>"""
        chapter_files.append((chapter_file, chapter_xhtml))
        manifest_items.append(f'<item id="{chapter_id}" href="{chapter_file}" media-type="application/xhtml+xml"/>')
        spine_items.append(f'<itemref idref="{chapter_id}"/>')
        nav_points.append(f'<navPoint id="navPoint-{idx}" playOrder="{play_order}"><navLabel><text>{html.escape(chapter_title)}</text></navLabel><content src="{chapter_file}"/></navPoint>')
        nav_links.append(f'<li><a href="{chapter_file}">{html.escape(chapter_title)}</a></li>')
        play_order += 1

    image_items = []
    image_id = 1
    added_images = set()
    for chap in chapters_data:
        for img_name in chap.get('images', []):
            if img_name in added_images:
                continue
            added_images.add(img_name)
            ext = os.path.splitext(img_name)[1].lower()
            media_type = {
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.png': 'image/png',
                '.gif': 'image/gif',
                '.webp': 'image/webp',
                '.svg': 'image/svg+xml'
            }.get(ext, 'image/jpeg')
            image_items.append((img_name, os.path.join(temp_dir, 'images', img_name), media_type))
            manifest_items.append(f'<item id="image_{image_id}" href="images/{img_name}" media-type="{media_type}"/>')
            image_id += 1

    content_opf = f"""<?xml version="1.0" encoding="utf-8"?>
<package xmlns="http://www.idpf.org/2007/opf" unique-identifier="BookId" version="2.0">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
    <dc:title>{html.escape(novel_data['title'])} - {html.escape(vol_info['title'])}</dc:title>
    <dc:creator>{html.escape(novel_data['author'])}</dc:creator>
    <dc:language>vi</dc:language>
    <dc:identifier id="BookId">urn:uuid:{book_id}</dc:identifier>
    <dc:description>{html.escape(Utils.summary_to_text(novel_data['summary']))}</dc:description>
    <dc:date>{created_at}</dc:date>
  </metadata>
  <manifest>{' '.join(manifest_items)}</manifest>
  <spine toc="ncx">{' '.join(spine_items)}</spine>
</package>"""

    toc_ncx = f"""<?xml version="1.0" encoding="utf-8"?>
<ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">
  <head>
    <meta name="dtb:uid" content="urn:uuid:{book_id}"/>
    <meta name="dtb:depth" content="1"/>
    <meta name="dtb:totalPageCount" content="0"/>
    <meta name="dtb:maxPageNumber" content="0"/>
  </head>
  <docTitle><text>{html.escape(novel_data['title'])} - {html.escape(vol_info['title'])}</text></docTitle>
  <navMap>{''.join(nav_points)}</navMap>
</ncx>"""

    nav_xhtml = f"""<?xml version="1.0" encoding="utf-8"?>
<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="vi">
  <head>
    <title>Mục lục</title>
    <link rel="stylesheet" type="text/css" href="styles/style.css"/>
  </head>
  <body>
    <h1>Mục lục</h1>
    <ol>{''.join(nav_links)}</ol>
  </body>
</html>"""

    with zipfile.ZipFile(epub_path, 'w') as epub:
        epub.writestr('mimetype', 'application/epub+zip', compress_type=zipfile.ZIP_STORED)
        epub.writestr('META-INF/container.xml', """<?xml version="1.0"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles>
    <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>
  </rootfiles>
</container>""")
        epub.writestr('OEBPS/content.opf', content_opf)
        epub.writestr('OEBPS/toc.ncx', toc_ncx)
        epub.writestr('OEBPS/nav.xhtml', nav_xhtml)
        epub.writestr('OEBPS/styles/style.css', STYLE_CSS)
        epub.writestr('OEBPS/text/intro.xhtml', intro_xhtml)
        for chapter_file, chapter_xhtml in chapter_files:
            epub.writestr(f'OEBPS/{chapter_file}', chapter_xhtml)
        for img_name, img_path, _media in image_items:
            if os.path.exists(img_path):
                epub.write(img_path, f'OEBPS/images/{img_name}')
    return epub_path


async def download_volumes(novel_data: Dict, selected_ids: List[int], export_formats: List[str], output_root: str, storage_state: Optional[Dict], progress_cb: Optional[Callable[[str], None]] = None) -> List[str]:
    progress_cb = progress_cb or (lambda _msg: None)
    crawler = HakoCrawler(storage_state=storage_state, progress_cb=progress_cb)
    export_formats = parse_export_formats(export_formats)
    volume_map = {vol['id']: vol for vol in novel_data['volumes']}
    selected_volumes = [volume_map[idx] for idx in selected_ids if idx in volume_map]
    if not selected_volumes:
        raise HakoError('Bạn chưa chọn tập hợp lệ để tải.')

    novel_dir = Path(output_root) / Utils.format_filename(novel_data['title'])
    novel_dir.mkdir(parents=True, exist_ok=True)
    created_files: List[str] = []

    for volume in selected_volumes:
        progress_cb(f"Bắt đầu tải tập: {volume['title']}")
        with tempfile.TemporaryDirectory() as temp_dir:
            playwright, browser, context, page = await crawler._new_page()
            try:
                connector = aiohttp.TCPConnector(ssl=False, limit=20)
                timeout = aiohttp.ClientTimeout(total=60)
                async with aiohttp.ClientSession(connector=connector, timeout=timeout) as session:
                    chapters_data = []
                    chapters = volume['chapters']
                    for idx, chapter in enumerate(chapters, start=1):
                        progress_cb(f"Tải chương {idx}/{len(chapters)}: {chapter['title']}")
                        html_content, imgs = await crawler.download_chapter(page, session, temp_dir, chapter, idx)
                        chapters_data.append({'title': chapter['title'], 'html': html_content, 'images': imgs})
                        if idx < len(chapters):
                            await asyncio.sleep(random.uniform(0.7, 1.5))

                if 'txt' in export_formats:
                    created_files.append(export_volume_txt(volume, novel_data, chapters_data, str(novel_dir)))
                if 'docx' in export_formats:
                    created_files.append(export_volume_docx(temp_dir, volume, novel_data, chapters_data, str(novel_dir)))
                if 'epub' in export_formats:
                    created_files.append(export_volume_epub(temp_dir, volume, novel_data, chapters_data, str(novel_dir)))
            finally:
                await context.close()
                await browser.close()
                await playwright.stop()
        progress_cb(f"Hoàn tất tập: {volume['title']}")

    archive_path = novel_dir / f"{Utils.format_filename(novel_data['title'])}-downloads.zip"
    with zipfile.ZipFile(archive_path, 'w', compression=zipfile.ZIP_DEFLATED) as archive:
        for file_path in created_files:
            archive.write(file_path, arcname=Path(file_path).name)
    created_files.append(str(archive_path))
    progress_cb(f"Đã đóng gói file ZIP: {archive_path.name}")
    return created_files


def reset_output_dir(output_root: str):
    path = Path(output_root)
    path.mkdir(parents=True, exist_ok=True)
    return str(path)
